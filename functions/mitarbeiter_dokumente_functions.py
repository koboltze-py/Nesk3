"""
Mitarbeiter-Dokumente – Hilfsfunktionen
Erstellen von Word-Dokumenten mit der DRK-Kopf-/Fußzeile aus der Vorlage
"""
import os
import shutil
from datetime import datetime
from pathlib import Path

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Pfad zur Vorlage mit Kopf-/Fußzeile
VORLAGE_PFAD = os.path.join(
    BASE_DIR, "Daten", "Mitarbeiter Vorlagen",
    "Kopf und Fußzeile",
    "Stärkemeldung 31.01.2026 bis 01.02.2026.docx",
)

# Ordner für erstellte Dokumente (kategorisiert)
DOKUMENTE_BASIS = os.path.join(BASE_DIR, "Daten", "Mitarbeiterdokumente")

KATEGORIEN = [
    "Stellungnahmen",
    "Bescheinigungen",
    "Dienstanweisungen",
    "Abmahnungen",
    "Lob & Anerkennung",
    "Sonstiges",
]


def sicherungsordner() -> str:
    """Gibt den Basispfad für Mitarbeiterdokumente zurück und legt ihn an."""
    os.makedirs(DOKUMENTE_BASIS, exist_ok=True)
    for kat in KATEGORIEN:
        os.makedirs(os.path.join(DOKUMENTE_BASIS, kat), exist_ok=True)
    return DOKUMENTE_BASIS


def lade_dokumente_nach_kategorie() -> dict[str, list[dict]]:
    """
    Gibt alle Dokumente je Kategorie zurück.
    Rückgabe: { "Stellungnahmen": [{"name": ..., "pfad": ..., "geaendert": ...}, ...], ... }
    """
    sicherungsordner()
    ergebnis: dict[str, list[dict]] = {}
    for kat in KATEGORIEN:
        ordner = os.path.join(DOKUMENTE_BASIS, kat)
        dateien = []
        if os.path.isdir(ordner):
            for fname in sorted(os.listdir(ordner)):
                if fname.lower().endswith((".docx", ".doc", ".pdf", ".txt")):
                    pfad = os.path.join(ordner, fname)
                    mtime = datetime.fromtimestamp(os.path.getmtime(pfad))
                    dateien.append({
                        "name": fname,
                        "pfad": pfad,
                        "geaendert": mtime.strftime("%d.%m.%Y %H:%M"),
                    })
        ergebnis[kat] = dateien
    return ergebnis


def erstelle_dokument_aus_vorlage(
    kategorie: str,
    titel: str,
    mitarbeiter: str,
    datum: str,
    inhalt: str,
    dateiname: str = "",
) -> str:
    """
    Erstellt ein neues Word-Dokument mit der DRK-Kopf-/Fußzeile aus der Vorlage.
    Gibt den Pfad der erstellten Datei zurück.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen.")

    sicherungsordner()

    # Dateiname generieren
    if not dateiname:
        safe_titel = "".join(c for c in titel if c.isalnum() or c in " _-").strip()[:40]
        safe_ma = "".join(c for c in mitarbeiter if c.isalnum() or c in " _-").strip()[:20]
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        dateiname = f"{safe_titel}_{safe_ma}_{stamp}.docx"

    ziel_pfad = os.path.join(DOKUMENTE_BASIS, kategorie, dateiname)

    # Vorlage öffnen (enthält Kopf-/Fußzeile)
    if os.path.isfile(VORLAGE_PFAD):
        doc = Document(VORLAGE_PFAD)
        # Bestehende Absätze im Body entfernen
        for para in doc.paragraphs[:]:
            p = para._element
            p.getparent().remove(p)
        for table in doc.tables[:]:
            t = table._element
            t.getparent().remove(t)
    else:
        # Fallback: leeres Dokument
        doc = Document()

    # Inhalt einfügen
    # Titel
    titel_para = doc.add_paragraph()
    run = titel_para.add_run(titel)
    run.bold = True
    run.font.size = Pt(16)
    titel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Leerzeile

    # Meta-Informationen
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Mitarbeiter: ").bold = True
    meta_para.add_run(mitarbeiter)
    meta_para.add_run(f"   Datum: ").bold = True
    meta_para.add_run(datum)

    doc.add_paragraph()  # Leerzeile

    # Hauptinhalt
    for zeile in inhalt.split("\n"):
        doc.add_paragraph(zeile)

    # Datum + Unterschrift
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f"Köln/Bonn Flughafen, {datum}")
    doc.add_paragraph()
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Unterschrift")

    doc.save(ziel_pfad)
    return ziel_pfad


def oeffne_datei(pfad: str) -> None:
    """Öffnet eine Datei mit der zugehörigen Windows-Standardanwendung."""
    import subprocess
    subprocess.Popen(["start", "", pfad], shell=True)


def loesche_dokument(pfad: str) -> bool:
    """Löscht eine Datei, gibt True zurück wenn erfolgreich."""
    try:
        if os.path.isfile(pfad):
            os.remove(pfad)
            return True
    except Exception:
        pass
    return False


def umbenennen_dokument(alter_pfad: str, neuer_name: str) -> str:
    """Benennt ein Dokument um. Gibt den neuen Pfad zurück."""
    ordner = os.path.dirname(alter_pfad)
    neuer_pfad = os.path.join(ordner, neuer_name)
    os.rename(alter_pfad, neuer_pfad)
    return neuer_pfad
